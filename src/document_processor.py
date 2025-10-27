import os
import logging
import datetime
from pathlib import Path
from typing import List, Dict, Any
import hashlib

# Fix PyPDF import
try:
    import pypdf
except ImportError:
    try:
        import PyPDF2 as pypdf
    except ImportError:
        pypdf = None

from docx import Document
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles processing of various document formats for clinical protocols."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        
    def process_file(self, file_path: Path) -> List[LangchainDocument]:
        """Process a single file and return chunked documents."""
        try:
            # Extract text based on file type
            if file_path.suffix.lower() == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                text = self._extract_docx_text(file_path)
            elif file_path.suffix.lower() in ['.md', '.markdown']:
                text = self._extract_markdown_text(file_path)
            elif file_path.suffix.lower() == '.txt':
                text = self._extract_txt_text(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_path}")
                return []
            
            if not text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return []
            
            # Create metadata
            metadata = self._create_metadata(file_path, text)
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create LangchainDocument objects
            documents = []
            for i, chunk in enumerate(chunks):
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_index": i,
                    "chunk_id": f"{metadata['file_hash']}_{i}",
                    "total_chunks": len(chunks)
                })
                
                documents.append(LangchainDocument(
                    page_content=chunk,
                    metadata=chunk_metadata
                ))
            
            logger.info(f"Processed {file_path}: {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            return []
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if pypdf is None:
            raise ImportError("PyPDF is not installed. Install with: pip install pypdf")
        
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            logger.warning(f"No text extracted from page {page_num + 1} of {file_path}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1} of {file_path}: {str(e)}")
                        continue
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {str(e)}")
            raise
        
        return text
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {str(e)}")
            raise
    
    def _extract_markdown_text(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            # Convert markdown to plain text (removes formatting)
            html = markdown.markdown(md_content)
            # Simple HTML tag removal (for basic conversion)
            import re
            text = re.sub('<[^<]+?>', '', html)
            return text
        except Exception as e:
            logger.error(f"Error reading Markdown file {file_path}: {str(e)}")
            raise
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {str(e)}")
            raise
    
    def _create_metadata(self, file_path: Path, text: str) -> Dict[str, Any]:
        """Create metadata for the document."""
        # Create file hash for unique identification
        file_hash = hashlib.md5(text.encode()).hexdigest()
        
        # Get file modification time correctly
        file_stats = file_path.stat()
        processed_at = datetime.datetime.fromtimestamp(file_stats.st_mtime).isoformat()
        
        return {
            "source": str(file_path),
            "filename": file_path.name,
            "file_type": file_path.suffix.lower(),
            "file_size": file_stats.st_size,
            "file_hash": file_hash,
            "processed_at": processed_at,
            "text_length": len(text),
            "protocol_type": self._infer_protocol_type(file_path.name),
        }
    
    def _infer_protocol_type(self, filename: str) -> str:
        """Infer the type of clinical protocol from filename."""
        filename_lower = filename.lower()
        
        if any(term in filename_lower for term in ['menopause', 'hormone', 'hrt']):
            return 'menopause'
        elif any(term in filename_lower for term in ['weight', 'obesity', 'glp1', 'semaglutide']):
            return 'weight_management'
        elif any(term in filename_lower for term in ['diabetes', 'glucose', 'insulin']):
            return 'diabetes'
        elif any(term in filename_lower for term in ['thyroid', 'tsh', 't4']):
            return 'thyroid'
        elif any(term in filename_lower for term in ['sleep']):
            return 'sleep'
        else:
            return 'general'